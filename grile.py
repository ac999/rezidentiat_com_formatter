from bs4 import BeautifulSoup
from docx import Document
import sys

def extract_questions_from_html(html_content, document_name):
    soup = BeautifulSoup(html_content, 'lxml')
    
    # Create a new Word document
    doc = Document()
    
    # Add a title to the document
    
    doc.add_heading(f"{document_name} - Baza de date", level=1)
    # Add empty paragraph
    doc.add_paragraph()
    
    # Extract each question
    for question_section in soup.find_all('div', class_="tab-pane"):
        question_text = question_section.find('h2').get_text(strip=True)

        # Add the question text in bold
        p = doc.add_paragraph()
        p.add_run(question_text).bold = True
        # Add empty paragraph
        doc.add_paragraph()

        # Initialize the question dictionary
        question_data = {
            'question': question_text,
            'options': [],
            'correct_answers': []
        }

        options = []
        correct_answers = []

        # Find answer options related to this question
        for answer_section in question_section.find_all('fieldset'):
            for answer_div in answer_section.find_all('div'):  # Find each div that contains an answer
                option_letter = answer_div.find('span', attrs={'class': 'ansNotation'}).get_text(strip=True).replace('.','')  # Extract the option letter (A, B, C, etc.)
                option_text = answer_div.find('label').get_text(strip=True)  # Extract the option text
                
                # Check if the answer is correct by looking for the 'fa-check' class inside the <i> tag
                if answer_div.find('i', class_='fa-check'):
                    correct_answers.append(option_letter)  # Append the correct answer letter
                
                # Add the option to the options list
                options.append(f"{option_letter}. {option_text}")
        
        question_data['options'] = options
        question_data['correct_answers'] = correct_answers

        # Add options to the document
        for option in options:
            doc.add_paragraph(option)  # Add each option as a new paragraph
        # Add empty paragraph
        doc.add_paragraph()
        # Add correct answers to the document
        doc.add_paragraph(f"Raspuns: {', '.join(question_data['correct_answers'])}\n")  # Adding correct answers
        

    # Save the document
    doc.save(f"{document_name}.docx")
    print(f"Document saved as {document_name}.docx")

# Example usage:
# html_content = "<html>...</html>"  # Your HTML content goes here
# extract_questions_from_html(html_content)

if __name__ == '__main__':
    if len(sys.argv) != 2:
        print(f"Usage: python {sys.argv[0]} <file_path>")
        sys.exit(1)
        
    html_file = sys.argv[1];
    try:
        with open(html_file, 'r', encoding='utf-8') as file:
            html_content = file.read()

        # Extract questions and format the output
        document_name = html_file.split('.')[-2].replace('\\','')
        extract_questions_from_html(html_content, document_name)
        
    except Exception as e:
        print(e)
